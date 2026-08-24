import os
import re
import unicodedata

import docx
import fitz


class DocumentPipeline:
    """Extract and normalize documents without destroying semantic context."""

    _CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
    _PAGE_NUMBER = re.compile(r"(?im)^\s*(?:page\s+)?\d+(?:\s+of\s+\d+)?\s*$")

    def __init__(self):
        self.documents = {}

    def ingest_pdf(self, file_path):
        with fitz.open(file_path) as document:
            return "\n\n".join(page.get_text("text", sort=True) for page in document)

    def ingest_docx(self, file_path):
        document = docx.Document(file_path)
        blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    blocks.append(" | ".join(values))
        return "\n".join(blocks)

    def ingest_txt(self, file_path):
        with open(file_path, "r", encoding="utf-8", errors="replace") as file:
            return file.read()

    def ingest_document(self, file_path):
        extension = os.path.splitext(file_path)[1].lower()
        readers = {
            ".pdf": self.ingest_pdf,
            ".docx": self.ingest_docx,
            ".txt": self.ingest_txt,
        }
        if extension not in readers:
            raise ValueError(f"Unsupported file format: {extension or 'unknown'}")

        text = readers[extension](file_path)
        if not text or not text.strip():
            raise ValueError("The document does not contain extractable text")
        self.documents[file_path] = text
        return text

    def clean_text(self, text):
        """Normalize extraction noise while preserving names, numbers and sentences.

        Lowercasing, punctuation stripping, stop-word removal and lemmatization are
        intentionally avoided: they reduce summary quality and erase context needed
        by named-entity recognition.
        """
        if not isinstance(text, str) or not text.strip():
            raise ValueError("Text must be a non-empty string")

        text = unicodedata.normalize("NFKC", text)
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = text.replace("\u00ad", "")
        text = self._CONTROL_CHARS.sub(" ", text)
        text = re.sub(r"(?<=\w)-\s*\n\s*(?=\w)", "", text)
        text = self._PAGE_NUMBER.sub("", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r" *\n *", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def process_document(self, file_path):
        return self.clean_text(self.ingest_document(file_path))
